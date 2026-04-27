import os
import tempfile
from io import BytesIO
from pathlib import Path

from pypdf import PdfReader
import docx

from all2txt import ExtractorError, TextDecoder


def extract_text_from_file(content: bytes, ext: str) -> str:
    """
    Verilen dosya içeriğinden (bytes) metin ayıklar.
    Desteklenen uzantılar: pdf, docx, txt, doc.
    """
    ext = ext.lower()
    if ext == "pdf":
        return extract_text_from_pdf(content)
    elif ext == "docx":
        return extract_text_from_docx(content)
    elif ext == "txt":
        return extract_text_from_txt(content)
    elif ext == "doc":
        return extract_text_from_doc(content)
    else:
        raise ValueError(f"Desteklenmeyen dosya formatı: {ext}")


def extract_text_from_txt(content: bytes) -> str:
    """
    Düz metin dosyasını okur. BOM ve yaygın kodlamaları sırayla dener.
    """
    for encoding in ("utf-8-sig", "utf-8", "cp1254", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_text_from_doc(content: bytes) -> str:
    """
    Eski Word (.doc) OLE biçimi. all2txt: Word/LibreOffice/antiword/OLE/python-bytes sırası.
    """
    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".doc")
        os.close(fd)
        Path(tmp_path).write_bytes(content)
        decoder = TextDecoder(
            preferred_tools=[
                "word",
                "libreoffice",
                "antiword",
                "wvtext",
                "catdoc",
                "ole",
                "strings",
            ],
            encoding="utf-8",
            fallback_encodings=["cp1254", "cp1252", "latin-1"],
        )
        text = decoder.decode_file(Path(tmp_path))
        if not text or not text.strip():
            raise ValueError("DOC dosyasından metin çıkarılamadı.")
        return text
    except ExtractorError as e:
        raise ValueError(f"DOC işlenemedi: {e}") from e
    finally:
        if tmp_path:
            Path(tmp_path).unlink(missing_ok=True)


def extract_text_from_pdf(content: bytes) -> str:
    text = ""
    try:
        reader = PdfReader(BytesIO(content))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        print(f"DEBUG: PDF'den {len(text)} karakter metin çıkarıldı.")
        return text
    except Exception as e:
        print(f"DEBUG: PDF okunurken hata: {e}")
        return ""

def extract_text_from_docx(content: bytes) -> str:
    file_stream = BytesIO(content)
    try:
        doc = docx.Document(file_stream)
        full_text = []
        
        # 1. Paragrafları oku
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # 2. Tablo içeriklerini oku
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        # 3. Header/Footer gibi gizli yerleri de tara (Alternatif tarama)
        if not full_text:
            for block in doc.element.body.iter():
                if block.tag.endswith('t'): # Text node
                    if block.text and block.text.strip():
                        full_text.append(block.text.strip())

        return "\n".join(full_text)
    except Exception as e:
        raise ValueError(f"DOCX okunurken teknik hata: {str(e)}")